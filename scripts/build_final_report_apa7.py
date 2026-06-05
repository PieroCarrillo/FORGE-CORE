from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIGURES = DOCS / "figures"
FIGURES.mkdir(parents=True, exist_ok=True)

OUT = DOCS / "Informe_Final_FORGE_CORE_APA7_Detallado.docx"

NAVY = "121D68"
BLUE = "1F4E79"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
SOFT_GREEN = "EAF8F0"
SOFT_GOLD = "FFF4DF"
SOFT_PURPLE = "F6F1FF"
TEXT = "111111"
BORDER = "B7C7D6"


def font(size: int = 20, bold: bool = False, italic: bool = False):
    candidates = []
    if bold and italic:
        candidates = ["C:/Windows/Fonts/arialbi.ttf", "C:/Windows/Fonts/timesbi.ttf"]
    elif bold:
        candidates = ["C:/Windows/Fonts/arialbd.ttf", "C:/Windows/Fonts/timesbd.ttf"]
    elif italic:
        candidates = ["C:/Windows/Fonts/ariali.ttf", "C:/Windows/Fonts/timesi.ttf"]
    else:
        candidates = ["C:/Windows/Fonts/arial.ttf", "C:/Windows/Fonts/times.ttf"]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_center(draw: ImageDraw.ImageDraw, box, text: str, fill: str, fnt):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=6)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, fill=fill, font=fnt, align="center", spacing=6)


def rounded_box(draw, box, fill, outline="#8AA4BA", radius=20, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#2D6F9F", width=5):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        points = [(x2, y2), (x2 - 16, y2 - 9), (x2 - 16, y2 + 9)]
    else:
        points = [(x2, y2), (x2 + 16, y2 - 9), (x2 + 16, y2 + 9)]
    draw.polygon(points, fill=fill)


def make_usil_logo(path: Path):
    img = Image.new("RGBA", (1000, 1000), (255, 255, 255, 0))
    d = ImageDraw.Draw(img)
    d.rounded_rectangle((25, 25, 975, 975), radius=85, fill=f"#{NAVY}", outline="#071148", width=5)
    try:
        f_u = ImageFont.truetype("C:/Windows/Fonts/times.ttf", 250)
        f_s = ImageFont.truetype("C:/Windows/Fonts/times.ttf", 355)
        f_i = ImageFont.truetype("C:/Windows/Fonts/timesi.ttf", 290)
        f_l = ImageFont.truetype("C:/Windows/Fonts/times.ttf", 245)
    except Exception:
        f_u = f_s = f_i = f_l = ImageFont.load_default()
    d.text((160, 330), "U", fill="white", font=f_u)
    d.text((370, 240), "S", fill="white", font=f_s)
    d.text((555, 355), "I", fill="white", font=f_i)
    d.text((655, 330), "L", fill="white", font=f_l)
    d.arc((300, 440, 575, 725), start=130, end=455, fill="white", width=12)
    d.ellipse((438, 565, 510, 637), outline="white", width=12)
    img.save(path)


def make_architecture(path: Path):
    img = Image.new("RGB", (1800, 1050), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "Arquitectura de despliegue AWS", fill="#123047", font=font(46, True))
    d.text((80, 110), "Separacion entre servidor de aplicacion y servidor de base de datos", fill="#5A6772", font=font(25))

    boxes = {
        "user": (80, 360, 360, 560),
        "nginx": (500, 220, 850, 390),
        "react": (500, 440, 850, 610),
        "api": (500, 660, 850, 860),
        "db": (1050, 360, 1430, 610),
        "sp": (1510, 360, 1710, 610),
        "worker": (1050, 710, 1430, 900),
    }
    colors = {
        "user": "#F5F8FB",
        "nginx": "#EAF7FF",
        "react": "#F4F8FB",
        "api": "#EAF2F8",
        "db": "#EAF8F0",
        "sp": "#FFF4DF",
        "worker": "#F6F1FF",
    }
    labels = {
        "user": "Usuario\nNavegador",
        "nginx": "EC2 App Server\nNginx puerto 80",
        "react": "Frontend\nReact build estatico",
        "api": "Backend\nNode.js + Express :4000",
        "db": "EC2 Database Server\nMariaDB :3306\nIP privada",
        "sp": "Stored\nprocedures",
        "worker": "Worker de metricas\n/proc + os + df",
    }
    for key, box in boxes.items():
        rounded_box(d, box, colors[key], "#95AFC4", 22, 3)
        draw_center(d, box, labels[key], "#123047", font(27, True))

    arrow(d, (360, 460), (500, 305))
    arrow(d, (360, 460), (500, 525))
    arrow(d, (850, 760), (1050, 485))
    arrow(d, (1430, 485), (1510, 485))
    arrow(d, (850, 760), (1050, 805))
    d.text((885, 450), "Conexion privada VPC", fill="#1F4E79", font=font(22, True))
    d.text((1015, 635), "Puerto 3306 permitido solo desde forge-app-sg", fill="#5A6772", font=font(20))
    img.save(path)


def make_er(path: Path):
    img = Image.new("RGB", (1800, 1150), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "Modelo logico de datos", fill="#123047", font=font(46, True))
    d.text((80, 110), "Entidades principales y relaciones de la base MariaDB forge_core", fill="#5A6772", font=font(25))
    entities = {
        "categories": (90, 250, 390, 420),
        "brands": (90, 520, 390, 690),
        "products": (560, 380, 920, 650),
        "inventory_movements": (1090, 190, 1510, 390),
        "order_items": (1090, 480, 1510, 700),
        "orders": (1090, 790, 1510, 990),
        "system_metrics": (90, 800, 390, 990),
        "buyer_messages": (560, 800, 920, 990),
    }
    content = {
        "categories": "categories\nid, name, slug",
        "brands": "brands\nid, name",
        "products": "products\nid, category_id, brand_id\nprice, stock, image_url",
        "inventory_movements": "inventory_movements\nproduct_id, type\nquantity, created_at",
        "order_items": "order_items\norder_id, product_id\nquantity, unit_price",
        "orders": "orders\ncustomer_name, total\nstatus, created_at",
        "system_metrics": "system_metrics\ncpu, memory, disk\nprocess_count",
        "buyer_messages": "buyer_messages\nproduct_id, author\nrating, body",
    }
    for key, box in entities.items():
        fill = "#EAF2F8" if key == "products" else "#F5F8FB"
        rounded_box(d, box, fill, "#95AFC4", 18, 3)
        draw_center(d, box, content[key], "#123047", font(23, True if key == "products" else False))
    arrow(d, (390, 335), (560, 465))
    arrow(d, (390, 605), (560, 535))
    arrow(d, (920, 480), (1090, 290))
    arrow(d, (920, 535), (1090, 590))
    arrow(d, (1510, 590), (1510, 890))
    arrow(d, (920, 895), (560, 895), fill="#6D7F90")
    d.text((980, 705), "orders 1:N order_items", fill="#5A6772", font=font(20))
    d.text((960, 245), "products 1:N inventory_movements", fill="#5A6772", font=font(20))
    img.save(path)


def make_checkout(path: Path):
    img = Image.new("RGB", (1800, 1000), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "Flujo de compra con pago simulado", fill="#123047", font=font(46, True))
    d.text((80, 110), "La transaccion protege el inventario ante compras simultaneas", fill="#5A6772", font=font(25))
    steps = [
        ("1", "Catalogo", "GET /api/products"),
        ("2", "Carrito", "POST /api/cart/quote"),
        ("3", "Checkout", "POST /api/orders/simulated-payment"),
        ("4", "Bloqueo", "SELECT ... FOR UPDATE"),
        ("5", "Pedido", "Descuenta stock y registra venta"),
    ]
    x = 80
    y = 290
    w = 285
    for i, (n, label, detail) in enumerate(steps):
        box = (x, y, x + w, y + 240)
        rounded_box(d, box, "#F4F8FB", "#95AFC4", 20, 3)
        d.ellipse((x + 25, y + 25, x + 85, y + 85), fill=f"#{BLUE}")
        draw_center(d, (x + 25, y + 25, x + 85, y + 85), n, "white", font(27, True))
        d.text((x + 30, y + 115), label, fill="#123047", font=font(27, True))
        d.multiline_text((x + 30, y + 158), detail, fill="#5A6772", font=font(20), spacing=5)
        if i < len(steps) - 1:
            arrow(d, (x + w, y + 120), (x + w + 55, y + 120))
        x += w + 65
    rounded_box(d, (250, 660, 1550, 860), "#FFF8E8", "#E0B24A", 18, 3)
    d.text((290, 700), "Punto critico de sincronizacion", fill="#7A5A00", font=font(29, True))
    d.multiline_text(
        (290, 750),
        "El stored procedure valida stock dentro de una transaccion. El bloqueo de filas evita que dos usuarios descuenten el mismo producto al mismo tiempo.",
        fill="#5A503C",
        font=font(22),
        spacing=7,
    )
    img.save(path)


def make_deployment(path: Path):
    img = Image.new("RGB", (1800, 950), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "Proceso de despliegue en AWS", fill="#123047", font=font(46, True))
    steps = [
        ("PC local", "Empaquetar proyecto\ntar.gz"),
        ("App Server", "Copiar codigo\n/opt/forge-core"),
        ("Backend", "npm install\nnpm run build\nsystemd restart"),
        ("Frontend", "npm install\nnpm run build\n/var/www"),
        ("Nginx", "Proxy /api\nservir SPA"),
    ]
    x = 100
    y = 280
    for label, detail in steps:
        box = (x, y, x + 280, y + 230)
        rounded_box(d, box, "#EAF2F8", "#95AFC4", 18, 3)
        d.text((x + 28, y + 45), label, fill="#123047", font=font(28, True))
        d.multiline_text((x + 28, y + 100), detail, fill="#5A6772", font=font(22), spacing=7)
        if label != "Nginx":
            arrow(d, (x + 280, y + 115), (x + 350, y + 115))
        x += 350
    rounded_box(d, (360, 650, 1440, 820), "#EAF8F0", "#7FBF99", 18, 3)
    d.text((400, 690), "Verificacion final", fill="#123047", font=font(28, True))
    d.text((400, 735), "curl /api/health -> status ok, database online. Navegador -> http://IP-publica-del-App-Server", fill="#5A6772", font=font(22))
    img.save(path)


def make_metrics_chart(path: Path):
    img = Image.new("RGB", (1800, 900), "white")
    d = ImageDraw.Draw(img)
    d.text((80, 50), "Metricas recolectadas del servidor de aplicacion", fill="#123047", font=font(46, True))
    values = [("CPU", 8, "%", "#5ED7EB"), ("RAM", 41, "%", "#78D66B"), ("Disco", 44, "%", "#F7C348"), ("Procesos", 117, "", "#8AA4FF")]
    max_val = 120
    base = 680
    x = 180
    for label, val, suffix, color in values:
        h = int(460 * min(val, max_val) / max_val)
        d.rounded_rectangle((x, base - h, x + 200, base), radius=16, fill=color)
        d.text((x + 28, base + 30), label, fill="#123047", font=font(28, True))
        d.text((x + 45, base - h - 55), f"{val}{suffix}", fill="#123047", font=font(28, True))
        x += 385
    d.text((80, 790), "Los porcentajes varian con la carga real de la instancia EC2. CPU baja indica baja actividad del servidor.", fill="#5A6772", font=font(22))
    img.save(path)


def set_update_fields(doc: Document):
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int]):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def set_base_styles(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.5)
        section.footer_distance = Inches(0.5)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    title = styles["Title"]
    title.font.name = "Times New Roman"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    title.font.size = Pt(16)
    title.font.bold = True
    title.font.color.rgb = RGBColor.from_string(TEXT)

    h1 = styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(TEXT)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.line_spacing = 2.0
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(0)

    h2 = styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(TEXT)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.line_spacing = 2.0
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(0)

    h3 = styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor.from_string(TEXT)
    h3.paragraph_format.line_spacing = 2.0
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(0)


def paragraph(doc: Document, text: str = "", style: str | None = None, indent: bool = True, align=None):
    p = doc.add_paragraph(text, style=style)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    if indent and style is None:
        p.paragraph_format.first_line_indent = Inches(0.5)
    if align is not None:
        p.alignment = align
    return p


def heading(doc: Document, text: str, level: int = 1):
    return paragraph(doc, text, style=f"Heading {level}", indent=False)


def bullet(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        p.add_run(item)


def numbered(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(0)
        p.add_run(item)


def apa_table(doc: Document, number: int, title: str, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    cap = paragraph(doc, f"Tabla {number}", indent=False)
    cap.runs[0].font.bold = True
    title_p = paragraph(doc, title, indent=False)
    title_p.runs[0].italic = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], h, True)
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def apa_figure(doc: Document, number: int, title: str, image: Path, width: float = 6.4, note: str | None = None):
    cap = paragraph(doc, f"Figura {number}", indent=False)
    cap.runs[0].font.bold = True
    title_p = paragraph(doc, title, indent=False)
    title_p.runs[0].italic = True
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image), width=Inches(width))
    n = paragraph(doc, note or "Nota. Elaboracion propia a partir de la implementacion del proyecto.", indent=False)
    n.runs[0].italic = True


def add_toc(doc: Document):
    heading(doc, "Tabla de contenidos", 1)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "La tabla de contenidos se actualiza automaticamente al abrir el documento en Word."
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(sep)
    run._r.append(text)
    run._r.append(end)
    doc.add_page_break()


def add_cover(doc: Document, logo: Path):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    def cover_p(text="", size=12, bold=False, space_after=0, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        r.font.size = Pt(size)
        r.font.bold = bold
        return p

    cover_p("Año de la esperanza y el fortalecimiento de la democracia", 12, False, 38)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    p.add_run().add_picture(str(logo), width=Inches(2.35))
    cover_p("", 12, False, 22)
    cover_p("Facultad de Ingeniería e Inteligencia Artificial", 12, False, 20)
    cover_p("Sistemas Operativos", 12, False, 20)
    cover_p("AVANCE 1 Y 2", 12, False, 20)
    cover_p("Grupo 4", 12, False, 8)
    cover_p("Pérez Sánchez, Alexander Augusto (2311311)", 11, False, 2)
    cover_p("Carrillo Malla, Piero (2411325)", 11, False, 2)
    cover_p("Paucar Gutierrez, Gabriel Rodrigo (2212468)", 11, False, 2)
    cover_p("Bolaños Zavala, Diego Alonso (2320374)", 11, False, 36)
    cover_p("Lima - Perú", 12, False, 2)
    cover_p("2026 - 01", 12, False, 0)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    body = doc.sections[-1]
    body.top_margin = Inches(1)
    body.bottom_margin = Inches(1)
    body.left_margin = Inches(1)
    body.right_margin = Inches(1)
    body.header_distance = Inches(0.5)
    add_page_number(body.header.paragraphs[0])


def references(doc: Document):
    heading(doc, "Referencias", 1)
    refs = [
        "Amazon Web Services. (s. f.). Amazon EC2 security groups for your EC2 instances. Amazon Elastic Compute Cloud User Guide. Recuperado el 19 de mayo de 2026, de https://docs.aws.amazon.com/AWSEC2/latest/UserGuide/ec2-security-groups.html",
        "Express. (s. f.). Routing. Express.js. Recuperado el 19 de mayo de 2026, de https://expressjs.com/en/guide/routing.html",
        "MariaDB. (s. f.). MariaDB Server documentation. Recuperado el 19 de mayo de 2026, de https://mariadb.com/docs/server/",
        "NGINX. (s. f.). NGINX reverse proxy. Recuperado el 19 de mayo de 2026, de https://docs.nginx.com/nginx/admin-guide/web-server/reverse-proxy/",
        "Node.js. (s. f.). Worker threads. Node.js documentation. Recuperado el 19 de mayo de 2026, de https://nodejs.org/api/worker_threads.html",
        "Tailwind Labs. (s. f.). Install Tailwind CSS with Vite. Tailwind CSS documentation. Recuperado el 19 de mayo de 2026, de https://tailwindcss.com/docs/installation/using-vite",
        "Vite. (s. f.). Getting started. Vite documentation. Recuperado el 19 de mayo de 2026, de https://vite.dev/guide/",
    ]
    for ref in refs:
        p = paragraph(doc, ref, indent=False)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)


def build():
    logo = FIGURES / "usil_logo_recreado.png"
    arch = FIGURES / "final_arquitectura_aws.png"
    er = FIGURES / "final_modelo_datos.png"
    checkout = FIGURES / "final_checkout.png"
    deploy = FIGURES / "final_despliegue.png"
    metrics = FIGURES / "final_metricas.png"
    make_usil_logo(logo)
    make_architecture(arch)
    make_er(er)
    make_checkout(checkout)
    make_deployment(deploy)
    make_metrics_chart(metrics)

    doc = Document()
    set_base_styles(doc)
    set_update_fields(doc)
    add_cover(doc, logo)

    heading(doc, "Agradecimiento", 1)
    paragraph(doc, "Agradecemos al docente del curso de Sistemas Operativos por orientar el desarrollo del trabajo final y por proponer un enfoque practico que permite relacionar conceptos teoricos con una implementacion real en servidores Linux. Asimismo, se reconoce el aporte de cada integrante del grupo en el analisis, diseno, construccion, despliegue y validacion de la herramienta FORGE CORE.")
    doc.add_page_break()

    heading(doc, "Resumen", 1)
    paragraph(doc, "El presente informe describe el desarrollo e implementacion de FORGE CORE, una herramienta web de comercio electronico ficticio orientada a productos tecnologicos como tarjetas graficas, memorias RAM, procesadores, unidades SSD y fuentes de poder. El sistema fue construido con una arquitectura distribuida en dos servidores Amazon EC2: un servidor de aplicacion que ejecuta Nginx, React y Node.js, y un servidor de base de datos dedicado con MariaDB.")
    paragraph(doc, "La propuesta permite demostrar conceptos del curso de Sistemas Operativos mediante procesos administrados con systemd, ejecucion concurrente de un worker para metricas, lectura de informacion del sistema Linux, uso de transacciones para sincronizacion de inventario y configuracion de seguridad de red mediante security groups. La aplicacion incluye catalogo, detalle de producto, carrito, pago simulado, dashboard administrativo y una seccion de comunidad con comentarios persistentes en MariaDB.")
    paragraph(doc, "Palabras clave: e-commerce, AWS EC2, MariaDB, Node.js, React, Sistemas Operativos, Nginx, transacciones.")
    doc.add_page_break()

    add_toc(doc)

    heading(doc, "Introduccion", 1)
    paragraph(doc, "En la actualidad, gran parte de las aplicaciones empresariales se ejecutan sobre infraestructura cloud y dependen de la correcta configuracion de servidores, redes, procesos, servicios y bases de datos. Por ello, un proyecto de e-commerce permite integrar de forma natural diversos conceptos de Sistemas Operativos, ya que requiere manejar procesos en segundo plano, recursos del servidor, comunicacion entre maquinas, seguridad de puertos, persistencia y sincronizacion de datos.")
    paragraph(doc, "FORGE CORE surge como una tienda ficticia de componentes de PC con apariencia profesional y funcionalidad demostrable. Aunque el flujo de pago es simulado, el sistema realiza operaciones reales sobre la base de datos: consulta productos, calcula totales, valida inventario, registra pedidos, descuenta stock y guarda comentarios de compradores. Esta caracteristica permite que la aplicacion sirva como evidencia tecnica y academica.")
    paragraph(doc, "El informe se organiza siguiendo la estructura solicitada en la guia del curso. Primero se presentan las generalidades del trabajo, luego el marco teorico, despues el desarrollo tecnico y finalmente las pruebas, resultados, conclusiones, recomendaciones y referencias.")

    heading(doc, "Generalidades del trabajo", 1)
    heading(doc, "Problematica", 2)
    paragraph(doc, "Muchas demostraciones academicas de tiendas virtuales se limitan a interfaces visuales sin persistencia real, sin control de inventario y sin separacion entre servidor de aplicacion y servidor de base de datos. Esta situacion dificulta evidenciar conceptos de Sistemas Operativos, porque no se observa administracion de servicios, configuracion de red, concurrencia, monitoreo de recursos ni sincronizacion de datos.")
    paragraph(doc, "El problema abordado consiste en construir una herramienta que no solo se vea como una tienda profesional, sino que tambien opere sobre dos servidores Linux en AWS y mantenga datos reales en MariaDB. De esta manera, el profesor puede interactuar con la aplicacion, publicar comentarios, simular compras y comprobar que los cambios quedan almacenados.")

    heading(doc, "Justificacion", 2)
    paragraph(doc, "La eleccion de un e-commerce de hardware permite trabajar con una tematica clara y visualmente atractiva. Los productos tecnologicos poseen atributos comparables, precios, stock, categorias y marcas, por lo que son adecuados para modelar tablas relacionales. Ademas, el checkout exige una operacion critica: validar y descontar inventario sin inconsistencias.")
    paragraph(doc, "Desde la perspectiva del curso, la solucion justifica el uso de dos instancias EC2 porque permite separar responsabilidades. El App Server publica la web y ejecuta procesos de Node.js, mientras que el DB Server concentra MariaDB. Esta division ayuda a explicar seguridad de puertos, redes privadas, servicios Linux, transacciones y monitoreo.")

    heading(doc, "Objetivos", 2)
    heading(doc, "Objetivo general", 3)
    paragraph(doc, "Implementar una herramienta e-commerce ficticia de componentes tecnologicos desplegada en AWS con dos servidores, aplicando conceptos de Sistemas Operativos y garantizando persistencia de datos en MariaDB.")
    heading(doc, "Objetivos especificos", 3)
    bullet(doc, [
        "Construir una SPA profesional usando React, TypeScript, Tailwind CSS y Vite.",
        "Implementar una API REST con Node.js y Express para catalogo, carrito, pedidos, comunidad y administracion.",
        "Configurar MariaDB en un servidor independiente y conectar el backend mediante IP privada.",
        "Crear scripts SQL, datos semilla y procedimientos almacenados para las operaciones principales.",
        "Aplicar transacciones y bloqueo de filas para proteger el inventario en el checkout.",
        "Recolectar metricas reales del servidor de aplicacion y mostrarlas en un dashboard.",
        "Desplegar la solucion con Nginx, systemd y security groups."
    ])

    apa_table(doc, 1, "Alcance funcional del proyecto", ["Modulo", "Funcion principal", "Persistencia"], [
        ["Inicio", "Presenta hero, carrusel destacado y metricas rapidas.", "Consulta productos y metricas."],
        ["Productos", "Lista, filtra y busca componentes tecnologicos.", "Lee products, brands y categories."],
        ["Detalle", "Muestra informacion completa de un producto.", "Lee product por slug."],
        ["Carrito", "Calcula cantidades, subtotal, IGV, envio y total.", "Consume API de cotizacion."],
        ["Checkout", "Simula pago y registra pedido.", "Inserta orders, order_items e inventory_movements."],
        ["Comunidad", "Permite comentarios de compradores.", "Inserta y lista buyer_messages."],
        ["Admin", "Muestra ventas, stock y metricas del sistema.", "Consulta orders, products y system_metrics."]
    ], [1800, 4300, 3260])

    heading(doc, "Marco teorico", 1)
    heading(doc, "Aplicaciones web SPA", 2)
    paragraph(doc, "Una Single Page Application concentra la experiencia de usuario en una sola pagina que cambia vistas de forma dinamica. En FORGE CORE, React administra el estado visual del catalogo, carrito, detalle, comunidad y panel administrativo. Vite se utiliza como herramienta de desarrollo y construccion, mientras que Tailwind CSS permite componer estilos responsivos con clases utilitarias.")
    heading(doc, "API REST con Express", 2)
    paragraph(doc, "Express permite definir rutas HTTP mediante metodos como GET, POST y PUT, cada uno asociado a un controlador que responde a una peticion especifica (Express, s. f.). Este enfoque se aplico para exponer endpoints de productos, categorias, carrito, pedidos, comunidad y administracion.")
    heading(doc, "MariaDB y transacciones", 2)
    paragraph(doc, "MariaDB funciona como sistema gestor relacional para mantener datos estructurados. En el proyecto se utilizan tablas normalizadas, claves foraneas y stored procedures. Para el checkout se usa una transaccion con bloqueo de filas, de modo que el stock se valida antes de descontarse y se evita una condicion de carrera cuando dos usuarios intentan comprar al mismo tiempo.")
    heading(doc, "Security groups en AWS", 2)
    paragraph(doc, "En Amazon EC2, un security group actua como firewall virtual para controlar el trafico entrante y saliente de una instancia (Amazon Web Services, s. f.). En FORGE CORE esta configuracion impide que MariaDB se exponga publicamente, ya que el puerto 3306 solo se permite desde el security group del App Server.")
    heading(doc, "Procesos, hilos y metricas del sistema operativo", 2)
    paragraph(doc, "Node.js incorpora el modulo worker_threads para ejecutar codigo JavaScript en paralelo dentro de hilos de trabajo (Node.js, s. f.). En este proyecto, el worker recolecta metricas del servidor sin bloquear la API principal. Las mediciones incluyen uso de CPU, memoria, disco y cantidad de procesos.")
    heading(doc, "Nginx y systemd", 2)
    paragraph(doc, "Nginx cumple el rol de servidor web y reverse proxy: sirve los archivos estaticos del frontend y redirige las rutas /api al backend Express. systemd administra el ciclo de vida del backend, permitiendo iniciarlo con el sistema, reiniciarlo y consultar logs de servicio.")

    heading(doc, "Desarrollo", 1)
    heading(doc, "Identificacion de requerimientos", 2)
    apa_table(doc, 2, "Requerimientos funcionales", ["Codigo", "Requerimiento", "Criterio de aceptacion"], [
        ["RF-01", "Visualizar catalogo de productos tecnologicos.", "La vista Productos lista minimo 50 productos desde MariaDB."],
        ["RF-02", "Buscar y filtrar productos.", "El usuario filtra por categoria y busca por texto."],
        ["RF-03", "Consultar detalle de producto.", "La vista muestra marca, descripcion, precio, stock e imagen."],
        ["RF-04", "Gestionar carrito.", "El usuario agrega, incrementa, reduce y elimina productos."],
        ["RF-05", "Simular pago.", "El backend registra pedido y descuenta stock con stored procedure."],
        ["RF-06", "Publicar comentarios.", "Los mensajes de Comunidad persisten en buyer_messages."],
        ["RF-07", "Consultar dashboard admin.", "El panel muestra ventas, stock y metricas del servidor."]
    ], [1100, 4300, 3960])
    apa_table(doc, 3, "Requerimientos no funcionales", ["Codigo", "Requerimiento", "Implementacion"], [
        ["RNF-01", "Seguridad de red.", "3306 solo desde forge-app-sg; HTTP publico en App Server."],
        ["RNF-02", "Disponibilidad del backend.", "Servicio systemd forge-core-api."],
        ["RNF-03", "Rendimiento academico.", "Frontend compilado y servido estaticamente por Nginx."],
        ["RNF-04", "Mantenibilidad.", "Estructura separada en frontend, backend, database, deploy y docs."],
        ["RNF-05", "Observabilidad.", "Metricas almacenadas en system_metrics y consultables por API."],
        ["RNF-06", "Integridad de datos.", "Transacciones SQL y claves foraneas."]
    ], [1100, 3500, 4760])

    heading(doc, "Analisis de hardware", 2)
    paragraph(doc, "La infraestructura se implemento con dos instancias EC2 Ubuntu. Ambas se ubicaron en la misma VPC para permitir comunicacion privada. El uso de instancias t3.micro responde al objetivo academico y al enfoque de bajo costo.")
    apa_table(doc, 4, "Hardware e infraestructura cloud utilizada", ["Recurso", "App Server", "DB Server"], [
        ["Nombre", "forge-app-server", "forge-db-server"],
        ["Tipo EC2", "t3.micro", "t3.micro"],
        ["Sistema operativo", "Ubuntu Server", "Ubuntu Server"],
        ["Rol", "Nginx, frontend y backend Node.js", "MariaDB dedicado"],
        ["IP privada", "172.31.31.144", "172.31.29.30"],
        ["Puerto principal", "80 publico y 4000 interno", "3306 privado"],
        ["Administracion", "SSH desde IP autorizada", "SSH desde IP autorizada"]
    ], [2700, 3330, 3330])

    heading(doc, "Analisis de software", 2)
    apa_table(doc, 5, "Stack de software aplicado", ["Capa", "Herramienta", "Uso dentro del proyecto"], [
        ["Frontend", "React + TypeScript", "Construccion de la SPA y tipado de componentes."],
        ["Estilos", "Tailwind CSS", "Diseno responsive, tarjetas, botones, navbar y paneles."],
        ["Build", "Vite", "Servidor de desarrollo y compilacion de frontend."],
        ["Backend", "Node.js + Express", "API REST y logica de negocio."],
        ["Base de datos", "MariaDB", "Persistencia relacional y stored procedures."],
        ["Servidor web", "Nginx", "Sirve la SPA y redirige /api."],
        ["Servicio Linux", "systemd", "Mantiene activo el backend."],
        ["Sistema operativo", "Ubuntu Server", "Entorno de ejecucion en AWS."]
    ], [1900, 2600, 4860])

    heading(doc, "Solucion del caso", 2)
    heading(doc, "Analisis y diseno", 3)
    paragraph(doc, "La solucion separa frontend, backend y base de datos en capas. El usuario interactua con una SPA servida por Nginx. Cuando necesita datos, la SPA consume rutas /api. Nginx redirige esas rutas a Express, y Express consulta MariaDB mediante mysql2. La base de datos mantiene el estado del catalogo, pedidos, inventario, metricas y comentarios.")
    apa_figure(doc, 1, "Arquitectura cloud de FORGE CORE", arch)
    paragraph(doc, "El modelo de datos se diseno para mantener entidades independientes y relaciones claras. Products se relaciona con categories y brands; orders se relaciona con order_items; products tambien se vincula con inventory_movements y buyer_messages.")
    apa_figure(doc, 2, "Modelo logico de datos de la base forge_core", er)

    apa_table(doc, 6, "Tablas de base de datos", ["Tabla", "Descripcion"], [
        ["users", "Clientes o usuarios ficticios utilizados para pedidos."],
        ["categories", "Categorias de producto como GPU, RAM, CPU, SSD y Fuente."],
        ["brands", "Marcas de hardware."],
        ["products", "Catalogo principal con precio, stock, slug, descripcion e imagen."],
        ["carts y cart_items", "Soporte para carrito persistente en futuras ampliaciones."],
        ["orders y order_items", "Pedidos simulados y detalle de productos comprados."],
        ["inventory_movements", "Auditoria de descuentos y reposiciones."],
        ["system_metrics", "Historico de metricas del sistema operativo."],
        ["buyer_messages", "Comentarios de compradores en Comunidad."]
    ], [2600, 6760])

    heading(doc, "Implementacion", 3)
    paragraph(doc, "La implementacion se dividio en cuatro frentes: frontend, backend, base de datos y despliegue. Cada frente mantiene archivos propios para facilitar mantenimiento y evidenciar entregables.")
    apa_table(doc, 7, "Estructura del repositorio", ["Carpeta", "Contenido"], [
        ["frontend", "SPA React, componentes visuales, estilos Tailwind y assets de productos."],
        ["backend", "API Express, validaciones, conexion a MariaDB y worker de metricas."],
        ["database", "schema.sql, procedures.sql, seed.sql y migraciones."],
        ["deploy", "Configuracion de Nginx, systemd y scripts de instalacion."],
        ["docs", "Informe, manual, presentacion, guion de video y runbook."]
    ], [2200, 7160])
    apa_table(doc, 8, "Endpoints principales de la API", ["Metodo", "Ruta", "Funcion"], [
        ["GET", "/api/health", "Verifica estado del backend y conexion con MariaDB."],
        ["GET", "/api/categories", "Lista categorias disponibles."],
        ["GET", "/api/products", "Lista productos con busqueda y filtros."],
        ["GET", "/api/products/:slug", "Devuelve detalle de producto."],
        ["POST", "/api/cart/quote", "Calcula subtotal, IGV, envio y total."],
        ["POST", "/api/orders/simulated-payment", "Ejecuta pago simulado y crea pedido."],
        ["GET", "/api/community/messages", "Lista comentarios de compradores."],
        ["POST", "/api/community/messages", "Guarda comentario en MariaDB."],
        ["GET", "/api/admin/dashboard", "Devuelve resumen administrativo."],
        ["GET", "/api/admin/system-metrics", "Devuelve metricas del App Server."],
        ["POST/PUT", "/api/admin/products", "Gestiona productos e inventario."]
    ], [1200, 3300, 4860])

    paragraph(doc, "El flujo de compra concentra la operacion transaccional mas importante. La API recibe los items del carrito, valida el payload y delega a MariaDB la creacion del pedido. El stored procedure bloquea los productos involucrados, revisa stock, inserta el pedido, inserta items, descuenta unidades y registra movimientos.")
    apa_figure(doc, 3, "Flujo transaccional del checkout", checkout)

    apa_table(doc, 9, "Procedimientos almacenados", ["Procedimiento", "Responsabilidad"], [
        ["sp_create_simulated_order", "Valida stock, bloquea productos, crea pedido, descuenta inventario y registra movimientos."],
        ["sp_restock_product", "Aumenta stock y registra movimiento de reposicion."],
        ["sp_admin_dashboard_summary", "Calcula ventas, pedidos, stock bajo y resumen rapido."],
        ["sp_register_system_metric", "Guarda CPU, RAM, disco y procesos recolectados por el worker."]
    ], [3300, 6060])

    heading(doc, "Configuracion de AWS", 3)
    paragraph(doc, "La configuracion de AWS se realizo con dos security groups. El objetivo fue permitir acceso web al App Server y restringir la base de datos a comunicacion privada. Esta separacion responde a buenas practicas de seguridad, ya que evita exponer MariaDB directamente a internet.")
    apa_table(doc, 10, "Reglas de seguridad implementadas", ["Security group", "Puerto", "Origen", "Motivo"], [
        ["forge-app-sg", "80", "0.0.0.0/0", "Permitir acceso web publico a la demo."],
        ["forge-app-sg", "22", "IP del desarrollador", "Administracion SSH controlada."],
        ["forge-db-sg", "3306", "forge-app-sg", "Permitir conexion privada solo desde la API."],
        ["forge-db-sg", "22", "IP del desarrollador", "Administracion puntual de MariaDB."]
    ], [2400, 1200, 2500, 3260])

    paragraph(doc, "El despliegue se automatizo de forma manual controlada: se empaqueta el proyecto desde la PC local, se copia al App Server, se instalan dependencias, se compilan backend y frontend, se reinicia el servicio systemd y se recarga Nginx.")
    apa_figure(doc, 4, "Proceso de despliegue de cambios hacia AWS", deploy)

    heading(doc, "Monitoreo del sistema operativo", 3)
    paragraph(doc, "El dashboard administrativo incluye metricas reales del App Server. El worker recolecta CPU, RAM, disco y cantidad de procesos, luego guarda el resultado en MariaDB. Esto permite relacionar la interfaz de la aplicacion con informacion del sistema operativo subyacente.")
    apa_figure(doc, 5, "Metricas tecnicas registradas en el dashboard", metrics)
    apa_table(doc, 11, "Conceptos de Sistemas Operativos evidenciados", ["Concepto", "Evidencia en el proyecto"], [
        ["Procesos", "El backend corre como servicio systemd y puede consultarse con systemctl."],
        ["Hilos", "Worker thread para recolectar metricas sin bloquear la API principal."],
        ["Sistema de archivos", "Lectura de /proc y uso de archivos de configuracion en Linux."],
        ["Comandos del sistema", "Uso de df para estimar almacenamiento disponible y usado."],
        ["Memoria", "Medicion y registro de RAM usada."],
        ["CPU", "Calculo de uso del procesador del App Server."],
        ["Sincronizacion", "Transacciones MariaDB y SELECT ... FOR UPDATE para proteger stock."],
        ["Seguridad", "Puertos controlados con security groups y usuario MariaDB limitado."]
    ], [2500, 6860])

    heading(doc, "Pruebas", 1)
    paragraph(doc, "Las pruebas se enfocaron en validar funcionamiento, persistencia, integridad de datos y conectividad entre servidores. Se considero correcto el resultado cuando la interfaz, la API y la base de datos mostraron el mismo estado.")
    apa_table(doc, 12, "Matriz de pruebas funcionales", ["Caso", "Procedimiento", "Resultado esperado", "Estado"], [
        ["PF-01", "Abrir http://IP-publica-App.", "Carga la pantalla Inicio.", "Aprobado"],
        ["PF-02", "Consultar /api/health.", "status ok y database online.", "Aprobado"],
        ["PF-03", "Entrar a Productos.", "Lista productos desde MariaDB.", "Aprobado"],
        ["PF-04", "Buscar producto por texto.", "La grilla se filtra sin recargar pagina.", "Aprobado"],
        ["PF-05", "Agregar producto al carrito.", "Producto aparece en Carrito.", "Aprobado"],
        ["PF-06", "Simular pago.", "Pedido se crea y stock disminuye.", "Aprobado"],
        ["PF-07", "Publicar comentario.", "Comentario persiste luego de recargar.", "Aprobado"],
        ["PF-08", "Abrir Admin.", "Metricas y resumen se muestran.", "Aprobado"]
    ], [1000, 3500, 3500, 1360])
    apa_table(doc, 13, "Pruebas tecnicas de infraestructura", ["Caso", "Comando o validacion", "Resultado esperado"], [
        ["PT-01", "sudo systemctl status forge-core-api", "Servicio active/running."],
        ["PT-02", "sudo systemctl status nginx", "Servicio active/running."],
        ["PT-03", "mysql -h 172.31.29.30 -u forge_app -p forge_core", "Conexion exitosa desde App Server."],
        ["PT-04", "Probar puerto 3306 desde internet", "Conexion bloqueada por security group."],
        ["PT-05", "SELECT COUNT(*) FROM products", "Catalogo con mas de 50 productos."],
        ["PT-06", "SELECT COUNT(*) FROM buyer_messages", "Comentarios almacenados."],
        ["PT-07", "SELECT * FROM system_metrics ORDER BY recorded_at DESC", "Metricas historicas registradas."]
    ], [1100, 4300, 3960])

    heading(doc, "Resultados", 1)
    paragraph(doc, "El sistema quedo operativo en AWS y accesible desde la IP publica del App Server. La API respondio correctamente, la base de datos se mantuvo en linea y el frontend mostro productos con imagenes, filtros, carrito, comunidad y dashboard.")
    apa_table(doc, 14, "Resultados obtenidos", ["Resultado", "Descripcion"], [
        ["Catalogo ampliado", "Se insertaron mas de 50 productos en MariaDB para una demo robusta."],
        ["Persistencia real", "Pedidos, comentarios, metricas e inventario quedan registrados en tablas."],
        ["Interfaz profesional", "Se incorporo estilo gaming premium, navbar oscuro, carrusel e imagenes."],
        ["Comunidad funcional", "Los visitantes pueden publicar comentarios que se guardan en MariaDB."],
        ["Metricas reales", "El dashboard consulta CPU, RAM, disco y procesos del App Server."],
        ["Seguridad aplicada", "MariaDB no se expone publicamente y usa conexion privada desde backend."]
    ], [2600, 6760])

    home = ROOT / "outputs" / "proposed-home.png"
    community = ROOT / "outputs" / "proposed-community-persistent.png"
    if home.exists():
        apa_figure(doc, 6, "Pantalla de inicio de FORGE CORE", home, 6.4, "Nota. Captura local del frontend antes del despliegue final en AWS.")
    if community.exists():
        apa_figure(doc, 7, "Seccion Comunidad con mensajes persistentes", community, 6.4, "Nota. Captura local de la seccion de comentarios integrada con la API.")

    heading(doc, "Conclusiones", 1)
    numbered(doc, [
        "FORGE CORE cumple el objetivo de implementar un e-commerce ficticio funcional con frontend, backend y base de datos separados.",
        "La arquitectura de dos servidores permite demostrar una comunicacion real entre App Server y Database Server por IP privada dentro de AWS.",
        "El uso de stored procedures y transacciones protege la integridad del inventario durante el checkout simulado.",
        "El worker de metricas permite evidenciar procesos, hilos, memoria, CPU, disco y lectura de informacion del sistema operativo Linux.",
        "La seccion Comunidad demuestra persistencia interactiva porque los comentarios escritos desde el navegador quedan almacenados en MariaDB.",
        "La configuracion de Nginx, systemd y security groups permite presentar una demo cercana a un entorno de produccion academico."
    ])

    heading(doc, "Recomendaciones", 1)
    numbered(doc, [
        "Asignar una IP elastica o dominio antes de una exposicion final para evitar que la URL cambie al reiniciar la instancia EC2.",
        "Mantener MariaDB sin exposicion publica y conservar el puerto 3306 restringido al security group del App Server.",
        "Agregar autenticacion para el panel Admin si el proyecto se extendiera fuera del contexto academico.",
        "Automatizar el despliegue con scripts o GitHub Actions para reducir errores manuales.",
        "Crear backups periodicos de MariaDB antes de realizar migraciones.",
        "Ampliar las pruebas de concurrencia simulando dos compras simultaneas del mismo producto."
    ])

    references(doc)

    heading(doc, "Anexos", 1)
    heading(doc, "Anexo A. Comandos principales de ejecucion", 2)
    apa_table(doc, 15, "Comandos utiles para operar la demo", ["Accion", "Comando"], [
        ["Conectar App Server", 'ssh -i "$env:USERPROFILE\\Downloads\\forge-core-key.pem" ubuntu@IP_APP'],
        ["Conectar DB Server", 'ssh -i "$env:USERPROFILE\\Downloads\\forge-core-key.pem" ubuntu@IP_DB'],
        ["Ver backend", "sudo systemctl status forge-core-api --no-pager"],
        ["Ver Nginx", "sudo systemctl status nginx --no-pager"],
        ["Ver MariaDB", "sudo systemctl status mariadb --no-pager"],
        ["Probar API", "curl http://127.0.0.1/api/health"],
        ["Ver logs backend", "sudo journalctl -u forge-core-api -n 100 --no-pager"]
    ], [2600, 6760])

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

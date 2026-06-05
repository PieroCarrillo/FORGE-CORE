from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIGURES = DOCS / "figures"
FIGURES.mkdir(parents=True, exist_ok=True)

OUT_REPORT = DOCS / "Informe_Implementacion_Configuracion_FORGE_CORE_APA.docx"
OUT_MANUAL = DOCS / "Manual_Uso_Basico_FORGE_CORE_APA.docx"

ACCENT = "1F4E79"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
BORDER = "B7C7D6"
TEXT = "111111"


def font(size: int = 20, bold: bool = False):
    names = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for name in names:
        if Path(name).exists():
            return ImageFont.truetype(name, size)
    return ImageFont.load_default()


def draw_centered(draw: ImageDraw.ImageDraw, box, text, fill, fnt):
    x1, y1, x2, y2 = box
    bbox = draw.multiline_textbbox((0, 0), text, font=fnt, spacing=4)
    w = bbox[2] - bbox[0]
    h = bbox[3] - bbox[1]
    draw.multiline_text((x1 + (x2 - x1 - w) / 2, y1 + (y2 - y1 - h) / 2), text, fill=fill, font=fnt, align="center", spacing=4)


def rounded_box(draw, box, fill, outline="#8AA4BA", radius=20, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def arrow(draw, start, end, fill="#2D6F9F", width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    if x2 >= x1:
        points = [(x2, y2), (x2 - 14, y2 - 8), (x2 - 14, y2 + 8)]
    else:
        points = [(x2, y2), (x2 + 14, y2 - 8), (x2 + 14, y2 + 8)]
    draw.polygon(points, fill=fill)


def make_architecture_diagram(path: Path):
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    title = font(42, True)
    h = font(28, True)
    body = font(22)
    small = font(18)
    d.text((80, 50), "Arquitectura AWS de FORGE CORE", fill="#123047", font=title)
    d.text((80, 105), "Usuario -> EC2 App Server -> EC2 Database Server", fill="#5A6772", font=body)

    boxes = {
        "user": (80, 280, 330, 470),
        "nginx": (450, 180, 760, 340),
        "api": (450, 420, 760, 610),
        "worker": (450, 660, 760, 820),
        "db": (920, 330, 1270, 560),
        "sp": (1340, 330, 1520, 560),
    }
    rounded_box(d, boxes["user"], "#F5F8FB")
    rounded_box(d, boxes["nginx"], "#EAF7FF")
    rounded_box(d, boxes["api"], "#EAF2F8")
    rounded_box(d, boxes["worker"], "#F6F1FF")
    rounded_box(d, boxes["db"], "#EAF8F0")
    rounded_box(d, boxes["sp"], "#FFF4DF")

    draw_centered(d, boxes["user"], "Usuario\nNavegador web", "#123047", h)
    draw_centered(d, boxes["nginx"], "EC2 App Server\nNginx + React build", "#123047", h)
    draw_centered(d, boxes["api"], "Backend Node.js\nExpress API :4000", "#123047", h)
    draw_centered(d, boxes["worker"], "Worker de metricas\n/proc + os + df", "#123047", h)
    draw_centered(d, boxes["db"], "EC2 Database Server\nMariaDB :3306\nIP privada", "#123047", h)
    draw_centered(d, boxes["sp"], "Stored\nProcedures", "#123047", h)

    arrow(d, (330, 375), (450, 260))
    arrow(d, (330, 375), (450, 510))
    arrow(d, (760, 510), (920, 445))
    arrow(d, (1270, 445), (1340, 445))
    arrow(d, (605, 610), (605, 660))

    d.text((450, 145), "Puerto 80 publico", fill="#1F4E79", font=small)
    d.text((855, 300), "Conexion privada VPC", fill="#1F4E79", font=small)
    d.text((915, 590), "3306 permitido solo desde el security group del App Server", fill="#5A6772", font=small)
    img.save(path)


def make_checkout_flow(path: Path):
    img = Image.new("RGB", (1600, 860), "white")
    d = ImageDraw.Draw(img)
    title = font(40, True)
    h = font(24, True)
    body = font(18)
    d.text((80, 50), "Flujo de compra con pago simulado e inventario", fill="#123047", font=title)

    steps = [
        ("1", "Catalogo", "GET /api/products"),
        ("2", "Carrito", "POST /api/cart/quote"),
        ("3", "Pago simulado", "POST /api/orders/simulated-payment"),
        ("4", "Transaccion SQL", "SELECT ... FOR UPDATE"),
        ("5", "Pedido", "Descuento de stock y registro"),
    ]
    x = 80
    y = 240
    w = 250
    for number, label, detail in steps:
        box = (x, y, x + w, y + 210)
        rounded_box(d, box, "#F4F8FB", "#95AFC4", 18, 3)
        d.ellipse((x + 22, y + 22, x + 76, y + 76), fill="#1F4E79")
        draw_centered(d, (x + 22, y + 22, x + 76, y + 76), number, "white", h)
        d.text((x + 28, y + 100), label, fill="#123047", font=h)
        d.multiline_text((x + 28, y + 140), detail, fill="#5A6772", font=body, spacing=4)
        if number != "5":
            arrow(d, (x + w, y + 105), (x + w + 70, y + 105))
        x += w + 90

    d.text((80, 560), "Punto critico de sincronizacion", fill="#123047", font=h)
    d.multiline_text(
        (80, 605),
        "La base de datos valida stock dentro de una transaccion. El bloqueo de filas evita que dos compras simultaneas descuenten el mismo inventario.",
        fill="#5A6772",
        font=body,
        spacing=8,
    )
    img.save(path)


def make_navigation_map(path: Path):
    img = Image.new("RGB", (1600, 900), "white")
    d = ImageDraw.Draw(img)
    title = font(40, True)
    h = font(24, True)
    body = font(18)
    d.text((80, 50), "Mapa basico de navegacion de la herramienta", fill="#123047", font=title)
    nodes = {
        "Inicio": (650, 90, 950, 220),
        "Productos": (190, 330, 490, 470),
        "Detalle": (650, 330, 950, 470),
        "Carrito": (1110, 330, 1410, 470),
        "Comunidad": (420, 650, 720, 790),
        "Admin": (880, 650, 1180, 790),
    }
    colors = {
        "Inicio": "#EAF2F8",
        "Productos": "#EAF8F0",
        "Detalle": "#FFF4DF",
        "Carrito": "#F6F1FF",
        "Comunidad": "#F5F8FB",
        "Admin": "#FBEFF1",
    }
    for name, box in nodes.items():
        rounded_box(d, box, colors[name])
        draw_centered(d, box, name, "#123047", h)
    arrow(d, (650, 155), (490, 400))
    arrow(d, (950, 155), (1110, 400))
    arrow(d, (490, 400), (650, 400))
    arrow(d, (950, 400), (1110, 400))
    arrow(d, (800, 220), (570, 650))
    arrow(d, (800, 220), (1030, 650))
    d.multiline_text(
        (80, 820),
        "El usuario puede iniciar en el hero, explorar productos, revisar detalle, comprar con pago simulado, comentar en Comunidad y consultar metricas en Admin.",
        fill="#5A6772",
        font=body,
    )
    img.save(path)


def make_metrics_chart(path: Path):
    img = Image.new("RGB", (1600, 760), "white")
    d = ImageDraw.Draw(img)
    title = font(40, True)
    h = font(24, True)
    body = font(18)
    d.text((80, 50), "Indicadores mostrados en el dashboard", fill="#123047", font=title)
    data = [("CPU", 8, "#5ED7EB"), ("RAM", 41, "#78D66B"), ("Disco", 44, "#F7C348"), ("Procesos", 117, "#8AA4FF")]
    x = 160
    max_h = 420
    for label, value, color in data:
        bar_h = int(max_h * (min(value, 120) / 120))
        d.rounded_rectangle((x, 620 - bar_h, x + 170, 620), radius=14, fill=color)
        d.text((x + 25, 635), label, fill="#123047", font=h)
        d.text((x + 35, 580 - bar_h), f"{value}%" if label != "Procesos" else str(value), fill="#123047", font=h)
        x += 330
    d.multiline_text(
        (80, 690),
        "Los valores pueden variar segun la carga real de la instancia EC2. CPU baja significa que el servidor esta casi sin uso.",
        fill="#5A6772",
        font=body,
    )
    img.save(path)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(10)
    run.font.bold = bold
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: Sequence[int]):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = OxmlElement("w:tblLayout")
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_pr.append(tbl_layout)
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[idx])


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_doc_defaults(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    section.different_first_page_header_footer = True
    add_page_number(section.header.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 2.0
    normal.paragraph_format.space_after = Pt(0)

    for name, size, color in [
        ("Title", 16, TEXT),
        ("Heading 1", 14, TEXT),
        ("Heading 2", 12, TEXT),
        ("Heading 3", 12, TEXT),
    ]:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(0)


def add_blank_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.0
    p.add_run("")
    doc.add_page_break()


def add_title_block(doc: Document, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(title)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(16)
    run.font.bold = True
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.line_spacing = 2.0
    r2 = p2.add_run(subtitle)
    r2.font.name = "Times New Roman"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    r2.font.size = Pt(12)
    doc.add_paragraph()


def p(doc: Document, text: str = "", style: str | None = None, first_indent: bool = True):
    paragraph = doc.add_paragraph(text, style=style)
    paragraph.paragraph_format.line_spacing = 2.0
    paragraph.paragraph_format.space_after = Pt(0)
    if first_indent and not style:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    return paragraph


def heading(doc: Document, text: str, level: int = 1):
    return p(doc, text, style=f"Heading {level}", first_indent=False)


def apa_table(doc: Document, number: int, title: str, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    cap = p(doc, f"Tabla {number}", first_indent=False)
    cap.runs[0].font.bold = True
    title_p = p(doc, title, first_indent=False)
    title_p.runs[0].italic = True
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for idx, head in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], head, True)
        shade_cell(table.rows[0].cells[idx], LIGHT_GRAY)
    for row_data in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row_data):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def apa_figure(doc: Document, number: int, title: str, image_path: Path, width_inches: float = 6.4):
    cap = p(doc, f"Figura {number}", first_indent=False)
    cap.runs[0].font.bold = True
    title_p = p(doc, title, first_indent=False)
    title_p.runs[0].italic = True
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.line_spacing = 1.0
    pic_p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    note = p(doc, "Nota. Elaboracion propia con base en la arquitectura implementada del proyecto.", first_indent=False)
    note.runs[0].italic = True
    return pic_p


def bullet(doc: Document, items: Iterable[str]):
    for item in items:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_after = Pt(0)
        para.add_run(item)


def numbered(doc: Document, items: Iterable[str]):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.line_spacing = 2.0
        para.paragraph_format.space_after = Pt(0)
        para.add_run(item)


def references(doc: Document, refs: Sequence[str]):
    heading(doc, "Referencias", 1)
    for ref in refs:
        para = p(doc, ref, first_indent=False)
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.first_line_indent = Inches(-0.5)


def build_report():
    arch = FIGURES / "figura_arquitectura_aws.png"
    checkout = FIGURES / "figura_flujo_checkout.png"
    metrics = FIGURES / "figura_metricas_dashboard.png"
    make_architecture_diagram(arch)
    make_checkout_flow(checkout)
    make_metrics_chart(metrics)

    doc = Document()
    set_doc_defaults(doc)
    add_blank_cover(doc)
    add_title_block(
        doc,
        "Informe de implementacion y configuracion de FORGE CORE",
        "E-commerce ficticio de hardware desplegado en AWS con React, Node.js, Express y MariaDB",
    )

    heading(doc, "Resumen", 1)
    p(doc, "El presente informe documenta el proceso de implementacion y configuracion de FORGE CORE, una herramienta web de comercio electronico ficticio orientada a la venta de componentes tecnologicos como tarjetas graficas, memorias RAM, procesadores, unidades SSD y fuentes de poder. La solucion se construyo con una arquitectura de dos servidores en Amazon EC2: un servidor de aplicacion con Nginx, frontend React y backend Node.js, y un servidor independiente para MariaDB. Esta separacion permite evidenciar una comunicacion cliente-servidor real, control de red mediante security groups, persistencia de datos, ejecucion de procedimientos almacenados y monitoreo de recursos del sistema operativo.")
    p(doc, "El proyecto cumple un alcance academico y funcional. Incluye catalogo, detalle de producto, carrito, pago simulado, dashboard administrativo, comunidad de comentarios y metricas del servidor. Asimismo, aplica conceptos de Sistemas Operativos como procesos, hilos, concurrencia, lectura de informacion del kernel Linux, transacciones y sincronizacion del inventario.")

    heading(doc, "Introduccion", 1)
    p(doc, "FORGE CORE fue disenado como una aplicacion demostrativa para integrar desarrollo web, administracion de base de datos y despliegue en la nube. La herramienta simula una tienda de componentes de PC, pero mantiene una estructura tecnica semejante a una aplicacion real: el frontend consume una API REST, la API se conecta a MariaDB por IP privada y las operaciones de compra se ejecutan dentro de procedimientos almacenados.")
    p(doc, "El despliegue se realizo en AWS usando instancias EC2 Ubuntu. El servidor de aplicacion publica la pagina por HTTP mediante Nginx y mantiene el backend con systemd. El servidor de base de datos no publica servicios web; MariaDB escucha por el puerto 3306 y acepta conexiones solo desde el servidor de aplicacion, de acuerdo con la configuracion de security groups recomendada para limitar el trafico entrante.")

    heading(doc, "Objetivos", 1)
    heading(doc, "Objetivo general", 2)
    p(doc, "Implementar y configurar una herramienta e-commerce ficticia de productos tecnologicos sobre infraestructura AWS, evidenciando conceptos de Sistemas Operativos mediante concurrencia, transacciones, monitoreo y separacion de servicios.")
    heading(doc, "Objetivos especificos", 2)
    bullet(doc, [
        "Desplegar una SPA en React, TypeScript y Tailwind CSS servida por Nginx.",
        "Implementar una API REST con Node.js y Express para catalogo, carrito, pedidos, comunidad y administracion.",
        "Configurar MariaDB en una instancia EC2 independiente y conectarla por red privada.",
        "Crear tablas, datos semilla y stored procedures para inventario, pedidos, metricas y comentarios.",
        "Registrar metricas de CPU, RAM, disco y procesos desde Linux para visualizarlas en el dashboard.",
        "Validar el flujo completo de compra simulada y persistencia de comentarios."
    ])

    heading(doc, "Arquitectura de la solucion", 1)
    p(doc, "La arquitectura implementada separa responsabilidades: la instancia de aplicacion atiende usuarios y peticiones HTTP; la instancia de base de datos concentra la persistencia. Esta division facilita explicar seguridad, redes, procesos y administracion de servicios Linux.")
    apa_figure(doc, 1, "Arquitectura general desplegada en AWS", arch)

    apa_table(doc, 1, "Componentes principales de la solucion", ["Componente", "Tecnologia", "Responsabilidad"], [
        ["Frontend", "React, TypeScript, Tailwind CSS, Vite", "Interfaz SPA con inicio, catalogo, detalle, carrito, comunidad y admin."],
        ["Backend", "Node.js, Express, mysql2", "API REST, validacion, calculo de carrito, pedidos simulados y consultas administrativas."],
        ["Base de datos", "MariaDB", "Persistencia de productos, pedidos, inventario, metricas y comentarios."],
        ["Proxy web", "Nginx", "Sirve el build estatico y redirige /api hacia Express."],
        ["Servicio Linux", "systemd", "Mantiene activo el backend y permite reinicio automatico."],
        ["Infraestructura", "AWS EC2 y Security Groups", "Separa servidor app y servidor BD dentro de la misma VPC."]
    ], [1900, 2300, 5160])

    heading(doc, "Implementacion paso a paso", 1)
    heading(doc, "Preparacion de AWS", 2)
    numbered(doc, [
        "Se creo o habilito la cuenta AWS y se trabajo en la region us-east-1.",
        "Se configuro un par de llaves SSH para conectarse desde Windows PowerShell a las instancias Ubuntu.",
        "Se crearon dos security groups: forge-app-sg para el servidor de aplicacion y forge-db-sg para el servidor de base de datos.",
        "El security group de aplicacion permite HTTP por el puerto 80 y SSH por el puerto 22 desde la IP autorizada del desarrollador.",
        "El security group de base de datos permite el puerto 3306 solo desde el security group del App Server."
    ])

    heading(doc, "Configuracion del servidor de base de datos", 2)
    p(doc, "En la instancia forge-db-server se instalo MariaDB, se habilito el servicio con systemd y se configuro el archivo de servidor para escuchar conexiones privadas. Luego se creo la base de datos forge_core y se cargaron los scripts schema.sql, procedures.sql, seed.sql y migraciones adicionales.")
    apa_table(doc, 2, "Tablas principales de MariaDB", ["Tabla", "Proposito"], [
        ["users", "Usuarios o clientes ficticios usados para pedidos."],
        ["categories", "Clasificacion de productos como GPU, RAM, CPU, SSD y Fuente."],
        ["brands", "Marcas asociadas a productos."],
        ["products", "Catalogo principal con precio, stock, descripcion e imagen."],
        ["carts y cart_items", "Estructura para carritos persistentes si se amplia el alcance."],
        ["orders y order_items", "Registro de pedidos simulados y sus lineas."],
        ["inventory_movements", "Auditoria de descuentos y reposiciones de stock."],
        ["system_metrics", "Historico de CPU, RAM, disco y procesos del App Server."],
        ["buyer_messages", "Comentarios de compradores guardados desde la seccion Comunidad."]
    ], [2500, 6860])

    heading(doc, "Procedimientos almacenados", 2)
    p(doc, "Los stored procedures concentran la logica critica de la base de datos. El caso mas importante es la creacion de pedidos simulados, porque valida stock, bloquea filas y descuenta inventario en una transaccion.")
    apa_table(doc, 3, "Stored procedures implementados", ["Procedimiento", "Funcion"], [
        ["sp_create_simulated_order", "Valida stock, ejecuta bloqueo con SELECT ... FOR UPDATE, inserta pedido, descuenta inventario y registra movimientos."],
        ["sp_restock_product", "Repone stock y registra movimiento administrativo."],
        ["sp_admin_dashboard_summary", "Devuelve resumen de ventas, pedidos, bajo stock y metricas."],
        ["sp_register_system_metric", "Guarda mediciones de CPU, RAM, disco y procesos."]
    ], [3300, 6060])

    heading(doc, "Configuracion del servidor de aplicacion", 2)
    p(doc, "En la instancia forge-app-server se instalo Node.js, Nginx y dependencias del proyecto. El backend se compilo con TypeScript y se ejecuto como servicio systemd. El frontend se compilo con Vite y el resultado se copio a /var/www/forge-core/frontend para que Nginx lo sirviera como sitio estatico.")
    apa_table(doc, 4, "Servicios del App Server", ["Servicio", "Comando de verificacion", "Resultado esperado"], [
        ["Backend Express", "sudo systemctl status forge-core-api --no-pager", "Servicio active/running."],
        ["Nginx", "sudo systemctl status nginx --no-pager", "Servicio active/running."],
        ["API local", "curl http://127.0.0.1:4000/api/health", "status ok y database online."],
        ["API publica via Nginx", "curl http://127.0.0.1/api/health", "status ok y database online."]
    ], [2100, 3900, 3360])

    heading(doc, "Configuracion de variables de entorno", 2)
    p(doc, "La conexion a la base de datos se definio en backend/.env. En produccion, DB_HOST debe usar la IP privada del DB Server, no la IP publica, para mantener la comunicacion dentro de la VPC.")
    apa_table(doc, 5, "Variables principales del backend", ["Variable", "Valor de produccion", "Descripcion"], [
        ["NODE_ENV", "production", "Modo de ejecucion del backend."],
        ["PORT", "4000", "Puerto interno de Express."],
        ["FRONTEND_ORIGIN", "http://34.230.29.124", "Origen permitido para CORS."],
        ["DB_HOST", "172.31.29.30", "IP privada del servidor MariaDB."],
        ["DB_PORT", "3306", "Puerto de MariaDB."],
        ["DB_NAME", "forge_core", "Base de datos de la aplicacion."]
    ], [2200, 2600, 4560])

    heading(doc, "Endpoints de la API", 1)
    apa_table(doc, 6, "Rutas REST disponibles", ["Metodo", "Ruta", "Uso"], [
        ["GET", "/api/health", "Verifica estado del backend y conexion con MariaDB."],
        ["GET", "/api/products", "Lista productos del catalogo con filtros y busqueda."],
        ["GET", "/api/products/:slug", "Obtiene detalle de un producto."],
        ["POST", "/api/cart/quote", "Calcula subtotal, IGV, envio y total."],
        ["POST", "/api/orders/simulated-payment", "Crea un pedido simulado y descuenta inventario."],
        ["GET", "/api/community/messages", "Lista comentarios guardados en MariaDB."],
        ["POST", "/api/community/messages", "Registra comentario de comprador."],
        ["GET", "/api/admin/dashboard", "Muestra resumen administrativo."],
        ["GET", "/api/admin/system-metrics", "Consulta metricas del sistema operativo."],
        ["POST/PUT", "/api/admin/products", "Gestiona productos e inventario."]
    ], [1300, 3300, 4760])

    heading(doc, "Flujo funcional de compra", 1)
    p(doc, "El flujo de compra usa un pago simulado. No existe integracion con pasarela real; sin embargo, el backend ejecuta validaciones y registra el pedido como si fuera una transaccion operativa.")
    apa_figure(doc, 2, "Secuencia de compra simulada y control de inventario", checkout)

    heading(doc, "Monitoreo y conceptos de Sistemas Operativos", 1)
    p(doc, "El backend incorpora un proceso secundario para recolectar metricas del servidor de aplicacion. La informacion se obtiene mediante modulos de Node.js, lectura del sistema Linux y comandos del sistema como df. Luego se persiste en MariaDB para ser consultada desde el dashboard.")
    apa_figure(doc, 3, "Indicadores del dashboard administrativo", metrics)
    apa_table(doc, 7, "Conceptos de Sistemas Operativos evidenciados", ["Concepto", "Implementacion en FORGE CORE"], [
        ["Procesos", "Backend Node.js ejecutado como servicio systemd."],
        ["Concurrencia", "Worker dedicado a recoleccion de metricas sin bloquear la API principal."],
        ["Kernel y sistema de archivos", "Lectura de /proc, uso del modulo os y comando df."],
        ["Sincronizacion", "Transacciones MariaDB y bloqueo SELECT ... FOR UPDATE durante checkout."],
        ["Seguridad", "Security groups, usuario MariaDB limitado y validacion de payloads."],
        ["Persistencia", "Registro durable de pedidos, productos, comentarios y metricas."]
    ], [2500, 6860])

    heading(doc, "Pruebas realizadas", 1)
    apa_table(doc, 8, "Plan de pruebas y resultados esperados", ["Prueba", "Procedimiento", "Resultado esperado"], [
        ["Conexion App -> BD", "mysql -h 172.31.29.30 -u forge_app -p forge_core", "Conexion exitosa solo desde App Server."],
        ["Catalogo", "Abrir /api/products y vista Productos", "Productos reales desde MariaDB."],
        ["Detalle", "Seleccionar Detalle en un producto", "Informacion, imagen, precio y stock visibles."],
        ["Carrito", "Agregar productos y cambiar cantidades", "Totales actualizados con IGV y envio."],
        ["Pago simulado", "Presionar Simular pago", "Pedido creado y stock descontado."],
        ["Stock insuficiente", "Comprar mas unidades que las disponibles", "Error controlado y rollback."],
        ["Comunidad", "Publicar comentario y recargar", "Comentario persiste en buyer_messages."],
        ["Metricas", "Abrir Admin o /api/admin/system-metrics", "CPU, RAM, disco y procesos consultables."]
    ], [2200, 3600, 3560])

    heading(doc, "Seguridad y red", 1)
    p(doc, "La base de datos no debe exponerse a internet. El puerto 3306 se restringe al security group del App Server, mientras que el usuario final solo accede por HTTP al servidor de aplicacion. Esta configuracion reduce el area de exposicion y permite demostrar separacion de capas.")
    apa_table(doc, 9, "Reglas de red aplicadas", ["Servidor", "Puerto", "Origen", "Justificacion"], [
        ["App Server", "80", "0.0.0.0/0", "Publicar la aplicacion web para la demo."],
        ["App Server", "22", "IP del desarrollador", "Administracion por SSH."],
        ["DB Server", "3306", "forge-app-sg", "Conexion privada desde backend a MariaDB."],
        ["DB Server", "22", "IP del desarrollador", "Administracion puntual de base de datos."]
    ], [2100, 1200, 2600, 3460])

    heading(doc, "Conclusiones", 1)
    p(doc, "La implementacion de FORGE CORE demuestra una arquitectura funcional de dos servidores con frontend, backend y base de datos separados. El uso de MariaDB en una instancia independiente, junto con una conexion por IP privada, permite explicar seguridad de red y aislamiento de servicios.")
    p(doc, "El proyecto tambien evidencia conceptos propios de Sistemas Operativos: el backend se administra como servicio, un worker recolecta metricas, la aplicacion interactua con informacion del sistema Linux y el checkout aplica sincronizacion transaccional para proteger el inventario.")
    p(doc, "Finalmente, la seccion Comunidad amplifica el alcance de la herramienta porque permite persistir interacciones de usuarios reales durante la demo. Esto confirma que la aplicacion no solo muestra datos fijos, sino que mantiene estado en la base de datos.")

    references(doc, [
        "Amazon Web Services. (s. f.). Security groups for your EC2 instances. Amazon EC2 User Guide. Recuperado el 19 de mayo de 2026, de https://docs.aws.amazon.com/AWSEC2/latest/UserGuide/ec2-security-groups.html",
        "Express. (s. f.). Express routing. Recuperado el 19 de mayo de 2026, de https://expressjs.com/en/guide/routing.html",
        "MariaDB. (s. f.). MariaDB documentation. Recuperado el 19 de mayo de 2026, de https://mariadb.com/docs/",
        "Node.js. (s. f.). Worker threads. Node.js documentation. Recuperado el 19 de mayo de 2026, de https://nodejs.org/api/worker_threads.html",
        "NGINX. (s. f.). NGINX reverse proxy. Recuperado el 19 de mayo de 2026, de https://docs.nginx.com/nginx/admin-guide/web-server/reverse-proxy",
        "Vite. (s. f.). Getting started. Recuperado el 19 de mayo de 2026, de https://vite.dev/guide/",
    ])

    doc.save(OUT_REPORT)


def build_manual():
    nav = FIGURES / "figura_mapa_navegacion.png"
    make_navigation_map(nav)
    home = ROOT / "outputs" / "proposed-home.png"
    community = ROOT / "outputs" / "proposed-community-persistent.png"

    doc = Document()
    set_doc_defaults(doc)
    add_blank_cover(doc)
    add_title_block(
        doc,
        "Manual de uso basico de FORGE CORE",
        "Guia operativa para navegar catalogo, carrito, comunidad y panel administrativo",
    )

    heading(doc, "Resumen", 1)
    p(doc, "Este manual explica el uso basico de FORGE CORE, una tienda ficticia de componentes tecnologicos. La herramienta permite explorar productos, consultar detalles, agregar articulos al carrito, ejecutar un pago simulado, publicar comentarios de compradores y revisar indicadores administrativos del servidor.")

    heading(doc, "Requisitos previos", 1)
    apa_table(doc, 1, "Requisitos para usar la herramienta", ["Elemento", "Requisito"], [
        ["Navegador", "Google Chrome, Microsoft Edge o navegador moderno compatible."],
        ["URL", "IP publica o dominio del App Server en AWS."],
        ["Conexion", "Acceso a internet para cargar la SPA y consumir la API."],
        ["Servidor app", "Nginx y forge-core-api deben estar activos."],
        ["Servidor BD", "MariaDB debe estar activo y accesible desde la IP privada del App Server."]
    ], [2700, 6660])

    heading(doc, "Acceso a la aplicacion", 1)
    numbered(doc, [
        "Abrir el navegador web.",
        "Ingresar la URL publica del App Server, por ejemplo http://34.230.29.124.",
        "Esperar a que cargue la pantalla de inicio de FORGE CORE.",
        "Si el navegador muestra una version antigua, presionar Ctrl + F5 para forzar recarga."
    ])
    if home.exists():
        apa_figure(doc, 1, "Pantalla de inicio con carrusel y metricas resumidas", home, 6.4)
    else:
        apa_figure(doc, 1, "Mapa de navegacion principal", nav, 6.4)

    heading(doc, "Mapa de navegacion", 1)
    p(doc, "La barra superior contiene los accesos principales: Inicio, Productos, Carrito, Comunidad y Admin. La navegacion no usa rutas separadas; cambia vistas internas dentro de la misma SPA.")
    apa_figure(doc, 2, "Rutas funcionales disponibles para el usuario", nav, 6.4)

    heading(doc, "Uso del Inicio", 1)
    p(doc, "La pantalla de Inicio presenta la identidad visual de FORGE CORE, un llamado a explorar componentes, metricas rapidas del servidor y un carrusel de productos destacados. Las flechas del carrusel permiten avanzar o retroceder entre productos principales.")
    apa_table(doc, 2, "Elementos visibles en Inicio", ["Elemento", "Funcion"], [
        ["CTA Explorar componentes", "Lleva al catalogo de productos."],
        ["Metricas CPU/RAM/Pedidos", "Muestra informacion resumida proveniente del dashboard."],
        ["Carrusel destacado", "Presenta productos destacados con imagen, marca y descripcion."],
        ["Boton Ver detalle", "Abre la vista detallada del producto seleccionado."]
    ], [3000, 6360])

    heading(doc, "Uso del catalogo de productos", 1)
    numbered(doc, [
        "Seleccionar la opcion Productos en el menu superior.",
        "Usar los filtros de categoria: Todos, GPU, RAM, CPU, SSD o Fuente.",
        "Escribir en el buscador para encontrar productos por nombre, marca o categoria.",
        "Presionar Detalle para revisar caracteristicas completas.",
        "Presionar Agregar para enviar el producto al carrito."
    ])
    apa_table(doc, 3, "Acciones principales del catalogo", ["Accion", "Resultado"], [
        ["Filtrar por categoria", "Reduce la lista a productos de una familia especifica."],
        ["Buscar por texto", "Muestra coincidencias por nombre, marca o descripcion."],
        ["Ver detalle", "Abre una vista con descripcion, precio, stock y especificaciones."],
        ["Agregar", "Suma una unidad al carrito y actualiza el conteo."]
    ], [3200, 6160])

    heading(doc, "Uso del detalle de producto", 1)
    p(doc, "La vista de detalle permite revisar la informacion completa antes de agregar un articulo al carrito. Es util para explicar que el catalogo no es una imagen estatica, sino informacion proveniente de MariaDB.")
    bullet(doc, [
        "Nombre del producto y marca.",
        "Categoria del componente.",
        "Precio y stock disponible.",
        "Descripcion comercial.",
        "Imagen del producto.",
        "Boton para agregar al carrito."
    ])

    heading(doc, "Uso del carrito y checkout", 1)
    numbered(doc, [
        "Entrar a Carrito desde la barra superior.",
        "Revisar los productos agregados.",
        "Aumentar o disminuir cantidades con los controles + y -.",
        "Eliminar productos con el icono de papelera.",
        "Ver subtotal, IGV, envio y total.",
        "Presionar Simular pago para crear el pedido."
    ])
    apa_table(doc, 4, "Estados posibles durante el checkout", ["Estado", "Significado", "Accion recomendada"], [
        ["Pago exitoso", "El pedido fue registrado y el stock se desconto en MariaDB.", "Mostrar confirmacion y revisar Admin."],
        ["Stock insuficiente", "El producto no tiene unidades suficientes.", "Reducir cantidad o elegir otro producto."],
        ["BD offline", "La API no puede conectarse a MariaDB.", "Revisar DB Server, security group y backend/.env."],
        ["Carrito vacio", "No hay productos seleccionados.", "Volver a Productos y agregar articulos."]
    ], [2200, 4300, 2860])

    heading(doc, "Uso de Comunidad", 1)
    p(doc, "La seccion Comunidad permite que compradores o visitantes de la demo dejen comentarios sobre productos. Estos mensajes se guardan en la tabla buyer_messages de MariaDB, por lo que siguen apareciendo despues de recargar la pagina.")
    numbered(doc, [
        "Entrar a Comunidad.",
        "Escribir el nombre del comprador.",
        "Seleccionar un producto del catalogo.",
        "Elegir una valoracion.",
        "Redactar el comentario.",
        "Presionar Publicar experiencia.",
        "Actualizar la pagina para comprobar que el comentario permanece guardado."
    ])
    if community.exists():
        apa_figure(doc, 3, "Seccion Comunidad con comentarios persistentes", community, 6.4)

    heading(doc, "Uso del panel Admin", 1)
    p(doc, "El panel Admin resume informacion operativa: ventas simuladas, pedidos, productos con stock bajo, estado de la base de datos y metricas del servidor. Esta seccion sirve para evidenciar tanto la administracion del e-commerce como conceptos de Sistemas Operativos.")
    apa_table(doc, 5, "Indicadores del panel Admin", ["Indicador", "Origen", "Interpretacion"], [
        ["Ventas simuladas", "orders y order_items", "Monto generado por pedidos ficticios."],
        ["Pedidos", "orders", "Cantidad de compras registradas."],
        ["Stock bajo", "products", "Productos que requieren reposicion."],
        ["CPU", "system_metrics", "Uso actual o reciente del procesador del App Server."],
        ["RAM", "system_metrics", "Porcentaje de memoria usada."],
        ["Disco", "system_metrics", "Uso del almacenamiento del servidor."],
        ["Procesos", "system_metrics", "Cantidad de procesos activos en Linux."]
    ], [2300, 2600, 4460])

    heading(doc, "Validaciones recomendadas durante una demo", 1)
    apa_table(doc, 6, "Checklist de demostracion", ["Paso", "Validacion", "Evidencia esperada"], [
        ["1", "Abrir la pagina publica.", "Carga la pantalla Inicio."],
        ["2", "Entrar a Productos.", "Se muestran productos con imagen y stock."],
        ["3", "Agregar producto al carrito.", "El carrito suma el producto."],
        ["4", "Simular pago.", "Se crea pedido y se descuenta inventario."],
        ["5", "Publicar comentario en Comunidad.", "El comentario persiste al recargar."],
        ["6", "Abrir Admin.", "Se ven pedidos y metricas del servidor."],
        ["7", "Abrir /api/health.", "Responde status ok y database online."]
    ], [900, 4000, 4460])

    heading(doc, "Solucion de problemas frecuentes", 1)
    apa_table(doc, 7, "Errores comunes y solucion", ["Problema", "Causa probable", "Solucion"], [
        ["No carga la pagina", "Instancia apagada, IP publica cambio o Nginx detenido.", "Revisar EC2, abrir la nueva IP y reiniciar Nginx."],
        ["API database offline", "DB Server apagado, puerto 3306 bloqueado o DB_HOST incorrecto.", "Encender DB Server, revisar security group y backend/.env."],
        ["No se guardan comentarios", "Tabla buyer_messages no existe o backend no fue actualizado.", "Ejecutar migracion 003 y reiniciar forge-core-api."],
        ["Imagenes no aparecen", "Archivos no copiados a /var/www o permisos incorrectos.", "Copiar assets y aplicar chmod 755/644."],
        ["CPU sale en 0%", "Servidor sin carga.", "Generar uso o explicar que la instancia esta inactiva."]
    ], [2500, 3400, 3460])

    heading(doc, "Cierre de uso", 1)
    p(doc, "Al finalizar una sesion de prueba o presentacion, se recomienda detener las instancias EC2 desde la consola de AWS si no se continuara usando la demo. Esto conserva los discos y la base de datos, pero evita mantener los servidores encendidos innecesariamente.")

    references(doc, [
        "Amazon Web Services. (s. f.). Security groups for your EC2 instances. Amazon EC2 User Guide. Recuperado el 19 de mayo de 2026, de https://docs.aws.amazon.com/AWSEC2/latest/UserGuide/ec2-security-groups.html",
        "Express. (s. f.). Express routing. Recuperado el 19 de mayo de 2026, de https://expressjs.com/en/guide/routing.html",
        "Node.js. (s. f.). Worker threads. Node.js documentation. Recuperado el 19 de mayo de 2026, de https://nodejs.org/api/worker_threads.html",
        "NGINX. (s. f.). NGINX reverse proxy. Recuperado el 19 de mayo de 2026, de https://docs.nginx.com/nginx/admin-guide/web-server/reverse-proxy",
    ])

    doc.save(OUT_MANUAL)


if __name__ == "__main__":
    build_report()
    build_manual()
    print(OUT_REPORT)
    print(OUT_MANUAL)
